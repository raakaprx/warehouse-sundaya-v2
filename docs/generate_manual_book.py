from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


@dataclass(frozen=True)
class Figure:
    image_path: Path
    caption: str


DOC_TITLE = "Sundaya Warehouse System V2.0"
DOC_SUBTITLE = "User Manual (Panduan Pengguna)"
DOC_CLASSIFICATION = "CONFIDENTIAL"
DOC_VERSION = "2.0"
DOC_DATE = date(2026, 4, 23)
DOC_ID = "SWS-UM-2.0"


def set_default_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Heading 1", 16, "8B0000"),
        ("Heading 2", 13, "1F1F1F"),
        ("Heading 3", 12, "1F1F1F"),
    ]:
        st = doc.styles[style_name]
        st.font.name = "Segoe UI"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)

    if "Caption" in doc.styles:
        cap = doc.styles["Caption"]
        cap.font.name = "Segoe UI"
        cap.font.size = Pt(9)
        cap.font.italic = True


def set_page_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    header = section.header
    header_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_p.text = f"{DOC_TITLE} | v{DOC_VERSION} | {DOC_CLASSIFICATION}"
    header_p.style = doc.styles["Normal"]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.runs[0].font.size = Pt(9)

    footer = section.footer
    footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(footer_p, prefix="Page ")


def add_page_number_field(paragraph, prefix: str = "") -> None:
    run = paragraph.add_run(prefix)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    # Placeholder text until fields are updated in Word
    run_text = OxmlElement("w:t")
    run_text.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(run_text)
    run._r.append(fld_end)


def add_toc(doc: Document) -> None:
    doc.add_heading("Daftar Isi", level=1)
    p = doc.add_paragraph(
        "Catatan: Daftar isi bersifat otomatis. Jika halaman belum muncul, buka dokumen di Microsoft Word lalu lakukan Update Field pada Daftar Isi."
    )
    p.runs[0].font.size = Pt(9)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o \"1-3\" \\h \\z \\u '

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    # Placeholder until TOC is updated in Word
    placeholder = OxmlElement("w:t")
    placeholder.text = "Klik kanan pada Daftar Isi, lalu pilih Update Field."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_cover(doc: Document, logo_path: Path) -> None:
    section = doc.sections[0]
    content_width = section.page_width - section.left_margin - section.right_margin

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(logo_path), width=Cm(8.0))

    doc.add_paragraph()

    title = doc.add_paragraph("USER MANUAL")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.runs[0]
    r.font.name = "Segoe UI Semibold"
    r.font.size = Pt(28)

    subtitle = doc.add_paragraph(DOC_SUBTITLE)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)

    system = doc.add_paragraph(DOC_TITLE)
    system.alignment = WD_ALIGN_PARAGRAPH.CENTER
    system.runs[0].font.size = Pt(16)
    system.runs[0].font.bold = True

    doc.add_paragraph()

    # Classification bar (table with shading)
    t = doc.add_table(rows=1, cols=1)
    t.autofit = True
    cell = t.cell(0, 0)
    cell.text = f"{DOC_CLASSIFICATION}"
    set_cell_shading(cell, "8B0000")
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string("FFFFFF")
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(12)

    doc.add_paragraph()

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    rows = [
        ("Dokumen ID", DOC_ID),
        ("Versi", DOC_VERSION),
        ("Tanggal", DOC_DATE.strftime("%d %B %Y")),
        ("Pemilik Dokumen", "PT Sundaya Indonesia"),
    ]
    for i, (k, v) in enumerate(rows):
        meta.cell(i, 0).text = k
        meta.cell(i, 1).text = v
        meta.cell(i, 0).paragraphs[0].runs[0].font.bold = True

    # Full-width spacing using an empty paragraph
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(8)
    spacer.paragraph_format.space_after = Pt(0)

    notice = doc.add_paragraph(
        "Dokumen ini berisi informasi internal PT Sundaya Indonesia dan ditujukan hanya untuk pihak yang berwenang. "
        "Dilarang mendistribusikan, menyalin, atau mengungkapkan isi dokumen tanpa persetujuan tertulis."
    )
    notice.runs[0].font.size = Pt(9)
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


def add_document_control(doc: Document) -> None:
    doc.add_heading("Document Control", level=1)

    ctrl = doc.add_table(rows=9, cols=2)
    ctrl.style = "Table Grid"
    rows = [
        ("Document Title", DOC_TITLE),
        ("Document Type", "User Manual"),
        ("Document ID", DOC_ID),
        ("Version", DOC_VERSION),
        ("Effective Date", DOC_DATE.strftime("%d %B %Y")),
        ("Classification", DOC_CLASSIFICATION),
        ("Prepared By", "[TBD]"),
        ("Reviewed By", "[TBD]"),
        ("Approved By", "[TBD]"),
    ]
    for i, (k, v) in enumerate(rows):
        ctrl.cell(i, 0).text = k
        ctrl.cell(i, 1).text = v
        ctrl.cell(i, 0).paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    doc.add_heading("Revision History", level=2)
    rev = doc.add_table(rows=2, cols=4)
    rev.style = "Table Grid"
    rev.rows[0].cells[0].text = "Version"
    rev.rows[0].cells[1].text = "Date"
    rev.rows[0].cells[2].text = "Description"
    rev.rows[0].cells[3].text = "Author"
    for c in rev.rows[0].cells:
        c.paragraphs[0].runs[0].font.bold = True
    rev.rows[1].cells[0].text = DOC_VERSION
    rev.rows[1].cells[1].text = DOC_DATE.strftime("%d %b %Y")
    rev.rows[1].cells[2].text = "Initial release untuk Sundaya Warehouse System V2.0"
    rev.rows[1].cells[3].text = "[TBD]"

    doc.add_paragraph()
    doc.add_heading("Distribution List", level=2)
    dist = doc.add_table(rows=5, cols=3)
    dist.style = "Table Grid"
    dist.rows[0].cells[0].text = "Role"
    dist.rows[0].cells[1].text = "Department"
    dist.rows[0].cells[2].text = "Purpose"
    for c in dist.rows[0].cells:
        c.paragraphs[0].runs[0].font.bold = True
    items = [
        ("OM", "Operational Maintenance", "Pengajuan permintaan material & receiving"),
        ("NOC", "Network Operation Center", "Review teknis, shipping, kontrol stok"),
        ("GM", "Management", "Approval final & monitoring eksekutif"),
        ("Programmer", "IT / Developer", "Kontrol sistem, audit, dan konfigurasi"),
    ]
    for i, (role, dept, purpose) in enumerate(items, start=1):
        dist.rows[i].cells[0].text = role
        dist.rows[i].cells[1].text = dept
        dist.rows[i].cells[2].text = purpose

    doc.add_page_break()


def add_paragraphs(doc: Document, lines: Iterable[str]) -> None:
    for line in lines:
        doc.add_paragraph(line)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_field_table(doc: Document, title: str, rows: list[tuple[str, str, str, str]]) -> None:
    doc.add_heading(title, level=3)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    headers = ["Field", "Deskripsi", "Aturan/Format", "Contoh"]
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    for field, desc, rule, example in rows:
        cells = t.add_row().cells
        cells[0].text = field
        cells[1].text = desc
        cells[2].text = rule
        cells[3].text = example


def add_figure(doc: Document, fig: Figure) -> None:
    section = doc.sections[0]
    content_width = section.page_width - section.left_margin - section.right_margin

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(fig.image_path), width=content_width)

    cap = doc.add_paragraph(fig.caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "Caption" in doc.styles:
        cap.style = doc.styles["Caption"]
    else:
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True


def add_role_summary_table(doc: Document) -> None:
    doc.add_heading("Ringkasan Role (Role-Based Access Control)", level=2)
    t = doc.add_table(rows=5, cols=4)
    t.style = "Table Grid"
    headers = ["Role", "Fokus", "Menu Utama", "Hak Akses Kritis"]
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    rows = [
        (
            "OM",
            "Operasional site & permintaan material",
            "Site Overview, Stock Master, Material Requests",
            "Buat request, konfirmasi receiving, input recycle & pemakaian",
        ),
        (
            "NOC",
            "Kontrol pusat, review teknis, shipping",
            "NOC Control, Review Request, Shipping Control",
            "Validasi request, input resi/ETA, respon system alerts",
        ),
        (
            "GM",
            "Monitoring eksekutif & approval final",
            "Executive Monitor, Review Request, Reports",
            "Final approval request, monitoring KPI, audit oversight",
        ),
        (
            "Programmer",
            "Administrasi sistem & kontrol penuh",
            "Programmer Control, Settings & Flow, Audit Logs",
            "Konfigurasi role/flow, audit, pemeliharaan sistem",
        ),
    ]
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            t.rows[r_i].cells[c_i].text = val


def add_chapter_1(doc: Document) -> None:
    doc.add_heading("Bab 1 - Pendahuluan", level=1)

    doc.add_heading("1.1 Tujuan Sistem", level=2)
    add_paragraphs(
        doc,
        [
            "Sundaya Warehouse System V2.0 adalah sistem manajemen material multi-site berbasis web untuk mendukung pengelolaan stok, permintaan material, proses persetujuan berlapis, pengiriman, penerimaan, serta mekanisme auto-alert stok kritis.",
            "Dokumen ini bertujuan memberikan panduan penggunaan sistem secara end-to-end sesuai peran (OM, NOC, GM, Programmer) dengan alur birokrasi standar perusahaan.",
        ],
    )

    doc.add_heading("1.2 Cakupan Site", level=2)
    add_paragraphs(
        doc,
        [
            "Sistem mencakup 3 area operasional (multi-site):",
            "1. Pusat (Warehouse Pusat)",
            "2. Papua (Site Operasional)",
            "3. Maluku (Site Operasional)",
        ],
    )

    doc.add_heading("1.3 Terminologi & Singkatan", level=2)
    t = doc.add_table(rows=6, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Istilah"
    t.rows[0].cells[1].text = "Definisi"
    for c in t.rows[0].cells:
        c.paragraphs[0].runs[0].font.bold = True
    items = [
        ("OM", "Operational Maintenance (pengaju permintaan material)"),
        ("NOC", "Network Operation Center (review teknis & shipping)"),
        ("GM", "General Manager / Executive Approver (approval final)"),
        ("Request ID", "Nomor permintaan (contoh: REQ-0020)"),
        ("Document No.", "Nomor dokumen permintaan (contoh: MR-20260423-4798)"),
    ]
    for i, (k, v) in enumerate(items, start=1):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v

    doc.add_heading("1.4 Persyaratan Teknis", level=2)
    add_paragraphs(
        doc,
        [
            "Perangkat & browser yang direkomendasikan:",
            "1. Google Chrome (versi terbaru) atau Microsoft Edge (versi terbaru)",
            "2. Resolusi layar minimal 1366x768 (disarankan 1920x1080)",
            "3. Koneksi internet stabil (disarankan minimum 5 Mbps)",
            "4. Akses URL aplikasi: [TBD] (Production/Staging)",
        ],
    )

    doc.add_heading("1.5 Keamanan & Kepatuhan", level=2)
    add_paragraphs(
        doc,
        [
            "Akses sistem diatur dengan Role-Based Access Control (RBAC). Setiap aktivitas penting (mis. perubahan stok, perubahan flow, approval) terekam dalam Audit Logs.",
            "Kebijakan password, masa berlaku sesi (session timeout), dan mekanisme reset password mengikuti standar keamanan internal perusahaan. Detail kebijakan: [TBD].",
        ],
    )

    doc.add_page_break()


def add_doc_conventions(doc: Document) -> None:
    doc.add_heading("Konvensi Dokumen", level=1)
    add_paragraphs(
        doc,
        [
            "Konvensi berikut digunakan di sepanjang dokumen ini:",
            '1. Tanda "[TBD]" berarti informasi belum tersedia dan perlu dikonfirmasi (mis. URL produksi, SLA, kontak support, atau detail kebijakan internal).',
            "2. Menu, tombol, dan label mengikuti tampilan aplikasi. Jika ada perbedaan minor (mis. perubahan teks tombol), ikuti label terbaru pada aplikasi.",
            "3. Screenshot pada dokumen ini adalah referensi visual. Untuk keamanan, data sensitif dapat disamarkan pada lingkungan produksi.",
        ],
    )
    doc.add_page_break()


def add_chapter_2(doc: Document, figures: dict[str, Figure]) -> None:
    doc.add_heading("Bab 2 - Memulai Sistem", level=1)

    doc.add_heading("2.1 Login (Langkah demi Langkah)", level=2)
    add_paragraphs(
        doc,
        [
            "1. Buka URL aplikasi Sundaya Warehouse System V2.0: [TBD].",
            "2. Masukkan Username dan Password sesuai akun yang diberikan.",
            "3. Klik tombol Masuk ke Sistem.",
            "4. Setelah login berhasil, Anda akan diarahkan ke Dashboard sesuai role.",
        ],
    )
    add_figure(doc, figures["login"])

    doc.add_heading("2.2 Navigasi Antarmuka", level=2)
    add_paragraphs(
        doc,
        [
            "Komponen utama antarmuka:",
            "1. Sidebar: menu utama sesuai role (OM/NOC/GM/Programmer).",
            "2. Area konten: menampilkan halaman/menu yang dipilih.",
            "3. Notifikasi: ikon di kanan atas untuk memantau notifikasi permintaan & alert stok.",
            "4. Settings: pengaturan profil akun dan keamanan.",
        ],
    )
    add_figure(doc, figures["om_site_overview"])

    add_role_summary_table(doc)

    doc.add_page_break()


def add_chapter_3(doc: Document) -> None:
    doc.add_heading("Bab 3 - Alur Birokrasi", level=1)

    doc.add_heading("3.1 Ringkasan Alur 5 Langkah", level=2)
    add_paragraphs(
        doc,
        [
            "Sistem menerapkan alur birokrasi 5 langkah berikut:",
            "Step 1: OM Request (OM mengajukan permintaan material)",
            "Step 2: NOC Technical Review (NOC melakukan review & validasi teknis)",
            "Step 3: GM Final Approval (GM memberikan persetujuan akhir)",
            "Step 4: NOC Shipping (NOC memproses pengiriman, input resi & ETA)",
            "Step 5: OM Receiving (OM melakukan konfirmasi penerimaan material)",
        ],
    )

    doc.add_heading("3.2 Tabel Alur Birokrasi", level=2)
    t = doc.add_table(rows=6, cols=5)
    t.style = "Table Grid"
    headers = ["Step", "Aktor", "Aktivitas Utama", "Output", "Notifikasi"]
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    rows = [
        (
            "1",
            "OM",
            "Buat permintaan (New Request) + tentukan urgensi & deadline",
            "Request ID & Document No. terbentuk",
            "Notifikasi ke NOC",
        ),
        (
            "2",
            "NOC",
            "Validasi ketersediaan stok, kelayakan, dan catatan teknis",
            "Status: Reviewed / Rejected [TBD]",
            "Notifikasi ke GM",
        ),
        (
            "3",
            "GM",
            "Approval final terhadap permintaan yang telah direview",
            "Status: Approved / Rejected",
            "Notifikasi ke NOC & OM",
        ),
        (
            "4",
            "NOC",
            "Proses pengiriman: input resi, ETA, bukti foto [TBD]",
            "Status: Shipping / Delivered",
            "Notifikasi ke OM",
        ),
        (
            "5",
            "OM",
            "Konfirmasi receiving: cek barang, input bukti terima [TBD]",
            "Status: Received / Fulfilled",
            "Notifikasi penyelesaian",
        ),
    ]
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            t.rows[r_i].cells[c_i].text = val

    doc.add_heading("3.3 Tabel Status (Standar)", level=2)
    st = doc.add_table(rows=1, cols=3)
    st.style = "Table Grid"
    st.rows[0].cells[0].text = "Status"
    st.rows[0].cells[1].text = "Makna"
    st.rows[0].cells[2].text = "Pemilik Aksi"
    for c in st.rows[0].cells:
        c.paragraphs[0].runs[0].font.bold = True

    statuses = [
        ("NEW", "Permintaan baru dibuat dan menunggu proses berikutnya", "OM"),
        ("MENUNGGU NOC REVIEW", "Menunggu validasi teknis", "NOC"),
        ("MENUNGGU GM APPROVAL", "Menunggu persetujuan akhir", "GM"),
        ("MENUNGGU NOC KIRIM", "Menunggu proses pengiriman", "NOC"),
        ("SHIPPING", "Sedang dalam pengiriman", "NOC"),
        ("RECEIVED", "Barang diterima oleh OM", "OM"),
        ("FULFILLED / SELESAI", "Permintaan selesai", "System"),
        ("REJECTED", "Permintaan ditolak", "NOC / GM"),
    ]
    for s in statuses:
        row = st.add_row().cells
        row[0].text, row[1].text, row[2].text = s

    doc.add_heading("3.4 Tingkat Urgensi", level=2)
    add_paragraphs(
        doc,
        [
            "Untuk menjaga prioritas operasional, permintaan dibagi dalam beberapa tingkat urgensi:",
            "1. Normal: pekerjaan rutin, deadline fleksibel.",
            "2. Penting: berdampak pada operasional, perlu diproses lebih cepat.",
            "3. Sangat Penting / Very Important: berdampak tinggi pada layanan/proyek.",
            "4. Critical: risiko penghentian operasional (contoh stok kritis).",
            "Kriteria detail penentuan urgensi dan SLA pemrosesan: [TBD].",
        ],
    )

    doc.add_page_break()


def add_chapter_4_om(doc: Document, figures: dict[str, Figure]) -> None:
    doc.add_heading("Bab 4 - OM (Operational Maintenance)", level=1)
    add_paragraphs(
        doc,
        [
            "Role OM berfokus pada pengajuan permintaan material untuk kebutuhan site, pemantauan status birokrasi, pelaporan recycle material, pencatatan pemakaian, serta konfirmasi penerimaan barang.",
            "Menu OM pada sidebar: Site Overview, Stock Master, Material Requests, Reports, Recycle Material, History Inventory, Notifications, Settings.",
        ],
    )

    doc.add_heading("4.1 Site Overview", level=2)
    add_paragraphs(
        doc,
        [
            "Halaman Site Overview menampilkan ringkasan permintaan material dan pergerakan terbaru untuk site yang dipilih.",
            "Fungsi utama:",
            "1. Monitoring total requests, in progress, completed.",
            "2. Melihat recent movements sebagai ringkasan transaksi terbaru.",
            "3. Aksi cepat untuk membuat permintaan material baru (jika tersedia).",
        ],
    )
    add_field_table(
        doc,
        "Definisi KPI (Site Overview)",
        [
            (
                "Total Requests",
                "Jumlah permintaan material pada scope dashboard.",
                "Bergantung filter/periode tampilan [TBD].",
                "20",
            ),
            (
                "In Progress",
                "Jumlah permintaan yang masih berjalan.",
                "Belum fulfilled/selesai.",
                "2",
            ),
            (
                "Completed",
                "Jumlah permintaan selesai.",
                "Status fulfilled/selesai.",
                "9",
            ),
            (
                "Recent Movements",
                "Ringkasan pergerakan terbaru.",
                "Menampilkan request ID, item, project, qty, status, waktu.",
                "[TBD]",
            ),
        ],
    )
    add_figure(doc, figures["om_site_overview"])

    doc.add_heading("4.2 Stock Master", level=2)
    add_paragraphs(
        doc,
        [
            "Stock Master menampilkan katalog material dan ketersediaan stok per lokasi.",
            "Langkah umum:",
            "1. Gunakan kolom pencarian untuk mencari berdasarkan SKU, nama item, atau spesifikasi.",
            "2. Gunakan filter lokasi untuk melihat stok per site.",
            "3. Klik Lihat Detail untuk melihat informasi material lebih lengkap.",
        ],
    )
    add_field_table(
        doc,
        "Elemen Data (Stock Master)",
        [
            ("SKU/SN", "Identitas unik material.", "Alfanumerik.", "7759"),
            (
                "Material Name",
                "Nama material.",
                "Teks.",
                "Positive Cable pole Ralis T5 3.8m",
            ),
            ("Category", "Kategori material.", "Teks.", "Cabling"),
            ("Current Stock", "Stok tersedia saat ini.", "Angka + unit.", "382 Units"),
            (
                "Warehouse Location",
                "Lokasi gudang/site asal stok.",
                "Pusat/Papua/Maluku.",
                "Pusat - Jakarta",
            ),
            (
                "Technical Specs",
                "Ringkasan spesifikasi.",
                "Opsional.",
                "3.8 meter positive cable",
            ),
        ],
    )
    add_figure(doc, figures["om_stock_master"])

    doc.add_heading("4.3 Material Requests", level=2)
    add_paragraphs(
        doc,
        [
            "Material Requests adalah modul utama OM untuk membuat dan memantau permintaan material sesuai alur birokrasi 5 langkah.",
            "Fungsi utama:",
            "1. Membuat permintaan baru (New Request).",
            "2. Memantau status dan progress step (OM Request -> NOC Review -> GM Approval -> NOC Shipping -> OM Receiving).",
            "3. Melihat tracking pengiriman (resi) dan ETA jika sudah diproses oleh NOC.",
            "4. Konfirmasi receiving dan melihat bukti terima (receipt) jika tersedia.",
        ],
    )
    add_field_table(
        doc,
        "Elemen Data (Ringkasan Request)",
        [
            ("Request ID", "Nomor permintaan.", "Otomatis.", "REQ-0020"),
            (
                "Document No.",
                "Nomor dokumen permintaan.",
                "Otomatis.",
                "MR-20260423-4798",
            ),
            ("Site Location", "Lokasi site pemohon.", "Pusat/Papua/Maluku.", "Maluku"),
            ("Project", "Nama proyek terkait.", "Teks.", "papua solar"),
            ("Destination", "Tujuan pengiriman.", "Teks.", "halsel bacan"),
            (
                "Urgency",
                "Tingkat urgensi.",
                "Normal/Penting/Sangat Penting/Critical.",
                "Very Important",
            ),
            ("Deadline", "Batas waktu kebutuhan.", "Tanggal.", "4/26/2026"),
            (
                "Tracking No.",
                "Nomor resi ekspedisi.",
                "Diisi NOC saat shipping.",
                "jne1245621331",
            ),
            ("ETA", "Perkiraan tanggal tiba.", "Diisi NOC.", "4/27/2026"),
            (
                "Receipt Confirmation",
                "Konfirmasi penerimaan.",
                "Diisi OM saat receiving [TBD].",
                "Items received",
            ),
        ],
    )
    add_figure(doc, figures["om_material_requests"])

    doc.add_heading("4.3.1 New Request (Form Permintaan)", level=3)
    add_paragraphs(
        doc,
        [
            "Langkah pembuatan permintaan:",
            "1. Klik tombol New Request pada halaman Material Requests.",
            "2. Isi lokasi site, nama proyek, destination, deadline kebutuhan, dan urgensi kebutuhan.",
            "3. Tambahkan material pada daftar material beserta quantity (Qty).",
            "4. Isi deskripsi/alasan kebutuhan material.",
            "5. Submit permintaan untuk diproses oleh NOC.",
            "Catatan: Nama tombol submit/finalisasi dapat berbeda sesuai versi aplikasi. Jika tidak muncul, hubungi support [TBD].",
        ],
    )
    add_field_table(
        doc,
        "Field Form New Request",
        [
            ("Lokasi Site", "Site pemohon.", "Wajib.", "Papua"),
            ("Nama Proyek", "Nama proyek/pekerjaan.", "Wajib.", "Papua Solar Phase 1"),
            ("Destination", "Tujuan pengiriman.", "Wajib.", "Site Papua"),
            ("Deadline Kebutuhan", "Tanggal kebutuhan.", "Wajib (mm/dd/yyyy).", "04/27/2026"),
            (
                "Urgensi Kebutuhan",
                "Prioritas permintaan.",
                "Penting/Sangat Penting (atau standar internal).",
                "Penting",
            ),
            ("Material List", "Daftar material dan qty.", "Minimal 1 item.", "Ralis 6 slot x5"),
            (
                "Description/Reason",
                "Alasan kebutuhan.",
                "Wajib.",
                "Untuk instalasi kabel proyek X",
            ),
        ],
    )
    add_figure(doc, figures["om_new_request_modal"])

    doc.add_heading("4.3.2 Konfirmasi Receiving (OM Receiving)", level=3)
    add_paragraphs(
        doc,
        [
            "Receiving adalah langkah birokrasi Step 5. OM melakukan verifikasi barang yang diterima dari NOC Shipping.",
            "Langkah umum (contoh):",
            "1. Buka request yang statusnya sudah sampai pada tahap pengiriman/diterima.",
            "2. Cocokkan jumlah dan jenis material dengan daftar permintaan.",
            "3. Jika tersedia, gunakan tombol View Receipt / Receipt Confirmation untuk melakukan konfirmasi penerimaan.",
            "4. Jika sistem meminta bukti terima (foto/tanda tangan), unggah sesuai prosedur perusahaan [TBD].",
            "Catatan: Detail form receiving dapat bervariasi sesuai implementasi. Jika menu receiving tidak muncul, hubungi NOC/Programmer.",
        ],
    )

    doc.add_heading("4.4 Reports", level=2)
    add_paragraphs(
        doc,
        [
            "Modul Reports menyediakan unduhan laporan dalam format PDF untuk kebutuhan audit dan monitoring.",
            "Jenis laporan umum:",
            "1. Stock Opname Report (stok semua material per site)",
            "2. Stock Mutation Report (mutasi stok masuk/keluar)",
            "3. Material Request Report (rekap permintaan material)",
            "4. Recent Movements Report (pergerakan material terbaru)",
            "5. Laporan Barang Diterima / Belum Diterima (jika tersedia)",
        ],
    )
    doc.add_heading("4.4.1 Jenis Laporan (Ringkas)", level=3)
    rep = doc.add_table(rows=1, cols=4)
    rep.style = "Table Grid"
    rep.rows[0].cells[0].text = "Report"
    rep.rows[0].cells[1].text = "Tujuan"
    rep.rows[0].cells[2].text = "Output"
    rep.rows[0].cells[3].text = "Catatan"
    for c in rep.rows[0].cells:
        c.paragraphs[0].runs[0].font.bold = True
    reports = [
        ("Stock Opname Report", "Rekap stok per site.", "PDF", "Gunakan untuk audit periodik."),
        ("Stock Mutation Report", "Mutasi stok masuk/keluar.", "PDF", "Validasi pergerakan stok."),
        ("Material Request Report", "Rekap request material.", "PDF", "Monitoring pipeline request."),
        ("Recent Movements Report", "Pergerakan terbaru.", "PDF", "Ringkasan aktivitas."),
        ("Laporan Barang Diterima", "Riwayat receiving.", "PDF", "Jika tersedia pada sistem."),
        ("Laporan Barang Belum Diterima", "Outstanding receiving.", "PDF", "Jika tersedia pada sistem."),
    ]
    for r in reports:
        cells = rep.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v
    add_figure(doc, figures["om_reports"])

    doc.add_heading("4.5 Recycle Material", level=2)
    add_paragraphs(
        doc,
        [
            "Recycle Material digunakan untuk pelaporan material rusak/terpakai agar dapat diproses (recycle) dan ditindaklanjuti oleh tim terkait.",
            "Praktik terbaik:",
            "1. Sertakan detail kondisi material dan deskripsi kerusakan.",
            "2. Lampirkan foto sebagai bukti (jika tersedia pada form).",
            "3. Pantau status tindak lanjut pada daftar riwayat recycle.",
        ],
    )
    add_field_table(
        doc,
        "Elemen Data (Recycle Material)",
        [
            ("Date", "Tanggal pelaporan.", "Wajib.", "23/04/2026"),
            (
                "Site/Reporter",
                "Site dan pelapor.",
                "Otomatis/diisi sistem [TBD].",
                "Papua / om",
            ),
            ("Material", "Material yang dilaporkan.", "Wajib.", "[TBD]"),
            ("Detail", "Rincian unit/serial.", "Opsional.", "[TBD]"),
            ("Photo", "Foto bukti.", "Opsional (disarankan).", "[TBD]"),
            (
                "Condition",
                "Kondisi material.",
                "Rusak/Terpakai/dll [TBD].",
                "Damaged",
            ),
            (
                "Description",
                "Deskripsi kerusakan/alasan.",
                "Wajib.",
                "Unit retak, perlu recycle",
            ),
            (
                "Status",
                "Status tindak lanjut.",
                "New/In Review/Resolved [TBD].",
                "New",
            ),
        ],
    )
    add_figure(doc, figures["om_recycle_material"])

    doc.add_heading("4.6 History Inventory", level=2)
    add_paragraphs(
        doc,
        [
            "History Inventory mencatat pemakaian material oleh OM termasuk tanggal, site, material, quantity, project, dan alasan penggunaan.",
            "Gunakan tombol Catat Pemakaian untuk menambahkan log pemakaian.",
        ],
    )
    add_field_table(
        doc,
        "Elemen Data (History Inventory)",
        [
            ("Tanggal", "Tanggal pemakaian.", "Wajib.", "23/04/2026"),
            ("Site", "Lokasi pemakaian.", "Wajib.", "Papua"),
            ("Material", "Nama material.", "Wajib.", "[TBD]"),
            ("Qty (pcs)", "Jumlah pemakaian.", "Wajib, angka.", "5"),
            ("Project", "Proyek terkait.", "Opsional/disarankan.", "Papua Solar"),
            ("Alasan", "Alasan pemakaian.", "Wajib.", "Pemeliharaan rutin"),
            ("Pelapor", "User pencatat.", "Otomatis.", "om"),
        ],
    )
    add_figure(doc, figures["om_history_inventory"])

    doc.add_heading("4.7 Notifications", level=2)
    add_paragraphs(
        doc,
        [
            "Notifications menampilkan riwayat notifikasi sistem terkait perubahan status permintaan dan alert stok.",
            "Fitur utama:",
            "1. Notification Feed (daftar notifikasi).",
            "2. Mark Read / Mark Read All untuk mengelola notifikasi.",
            "3. Monitor Notifikasi (panel ringkas) dari ikon notifikasi.",
        ],
    )
    add_field_table(
        doc,
        "Kategori Notifikasi (Contoh)",
        [
            (
                "Perubahan Status Request",
                "Notifikasi saat request berpindah step (review/approval/shipping/receiving).",
                "Otomatis.",
                "Permintaan papua solar telah direview NOC.",
            ),
            (
                "Informasi Shipping",
                "Notifikasi saat resi/ETA diinput atau status pengiriman berubah.",
                "Otomatis.",
                "Resi: jne1245621331",
            ),
            (
                "Alert Stok",
                "Notifikasi warning/critical saat stok mendekati threshold.",
                "Otomatis (berdasarkan threshold).",
                "Stock ... mendekati batas (18/15)",
            ),
            (
                "Sistem",
                "Notifikasi umum sistem (maintenance, perubahan akses) [TBD].",
                "Opsional.",
                "[TBD]",
            ),
        ],
    )
    add_figure(doc, figures["om_notifications"])
    add_figure(doc, figures["om_notifications_monitor"])

    doc.add_heading("4.8 Settings", level=2)
    add_paragraphs(
        doc,
        [
            "Settings digunakan untuk mengelola profil akun dan pengaturan keamanan.",
            "Ruang lingkup:",
            "1. Profile Akun (nama, email, nomor telepon, bahasa).",
            "2. Keamanan (password dan preferensi keamanan) [TBD].",
        ],
    )
    add_field_table(
        doc,
        "Field Profile Akun",
        [
            ("Username/Nama Lengkap", "Identitas akun.", "Tidak boleh kosong.", "om"),
            ("Email Address", "Email user.", "Format email valid.", "user@sundaya.co.id"),
            ("Nomor Telepon", "Kontak user.", "Format internasional disarankan.", "+62 812 3456 7890"),
            ("Bahasa/Language", "Preferensi bahasa UI.", "Dropdown.", "English"),
            ("Role Akun", "Role akses.", "Read-only (umumnya).", "OM"),
        ],
    )
    add_figure(doc, figures["om_settings"])

    doc.add_page_break()


def add_chapter_5_noc(doc: Document, figures: dict[str, Figure]) -> None:
    doc.add_heading("Bab 5 - NOC (Network Operation Center)", level=1)
    add_paragraphs(
        doc,
        [
            "Role NOC memiliki tanggung jawab utama pada kontrol stok pusat, review teknis permintaan, respon alert stok, serta proses shipping (pengiriman) material ke site.",
            "Menu NOC pada sidebar: NOC Control, Stock Master, Review Request, Reports, Recycle Material, History Inventory, Shipping Control, Notifications, System Alerts, Audit Logs, Settings.",
        ],
    )

    doc.add_heading("5.1 NOC Control Dashboard", level=2)
    add_paragraphs(
        doc,
        [
            "NOC Control Dashboard menampilkan ringkasan analitik (KPI) gudang untuk membantu pengambilan keputusan operasional.",
            "Contoh KPI: total sites, total stock items, critical needs, weekly turnover, serta grafik trend inbound/outbound.",
        ],
    )
    add_field_table(
        doc,
        "Definisi KPI (NOC Control Dashboard)",
        [
            ("Total Sites", "Jumlah site yang dimonitor.", "Tetap sesuai konfigurasi sistem.", "3"),
            (
                "Total Stock Items",
                "Jumlah item pada katalog/stock master.",
                "Bergantung data master.",
                "[TBD]",
            ),
            (
                "Critical Needs",
                "Jumlah kebutuhan kritis yang perlu tindakan.",
                "Berdasarkan alert/threshold.",
                "0",
            ),
            (
                "Weekly Turnover",
                "Perubahan aktivitas mingguan.",
                "Berdasarkan pergerakan inbound/outbound.",
                "+12%",
            ),
        ],
    )
    add_figure(doc, figures["noc_control"])

    doc.add_heading("5.2 Stock Master", level=2)
    add_paragraphs(
        doc,
        [
            "Stock Master pada role NOC memungkinkan pengelolaan katalog material (create/edit) termasuk threshold alert stok.",
            "Aksi umum:",
            "1. Add New Material untuk menambahkan material baru.",
            "2. Edit material untuk memperbarui spesifikasi, stok awal, threshold, foto, dan lokasi warehouse.",
        ],
    )
    add_field_table(
        doc,
        "Field Add/Edit Material (Central Inventory Catalog)",
        [
            ("SKU/SN", "Identitas material (SKU atau serial).", "Wajib, unik.", "SDY-BAT-12V"),
            ("Category", "Kategori material.", "Wajib.", "Battery"),
            ("Item Code", "Kode internal material.", "Wajib.", "SAP-7759"),
            ("Technical Specs", "Spesifikasi teknis ringkas.", "Opsional/disarankan.", "3.8 meter positive cable"),
            ("Material Name", "Nama material.", "Wajib.", "Positive Cable pole Ralis T5 3.8m"),
            ("Material Photo", "Foto material.", "JPG/PNG, maks 2MB (sesuai UI).", "[TBD]"),
            ("Initial Stock", "Stok awal saat input.", "Angka >= 0.", "382"),
            ("Threshold Alert", "Batas minimum stok untuk alert.", "Angka >= 0.", "37"),
            ("Warehouse Location", "Lokasi stok utama.", "Pusat/Papua/Maluku.", "Pusat"),
        ],
    )
    add_figure(doc, figures["noc_stock_master"])
    add_figure(doc, figures["noc_add_material_modal"])
    add_figure(doc, figures["noc_edit_material_modal"])

    doc.add_heading("5.3 Review Request (NOC Technical Review)", level=2)
    add_paragraphs(
        doc,
        [
            "Review Request adalah langkah birokrasi Step 2. NOC memvalidasi permintaan OM sebelum diteruskan ke GM untuk approval final.",
            "Checklist review (contoh):",
            "1. Ketersediaan stok di Pusat dan site terkait.",
            "2. Kesesuaian material dengan kebutuhan proyek.",
            "3. Alternatif material (jika stok terbatas) [TBD].",
            "4. Catatan teknis dan rekomendasi.",
        ],
    )
    add_field_table(
        doc,
        "Elemen Data (Review Request)",
        [
            ("Request ID", "Nomor permintaan.", "Read-only.", "REQ-0020"),
            ("Document No.", "Nomor dokumen.", "Read-only.", "MR-20260423-4798"),
            ("Site Location", "Lokasi site pemohon.", "Read-only.", "Maluku"),
            ("Material", "Material yang diminta.", "Read-only.", "Ralis 6 slot"),
            ("Total Items", "Jumlah jenis & unit.", "Read-only.", "1 Types - 5 Units"),
            ("Project", "Nama proyek.", "Read-only.", "papua solar"),
            ("Urgency", "Urgensi permintaan.", "Read-only.", "Very Important"),
            ("Deadline", "Batas waktu kebutuhan.", "Read-only.", "4/26/2026"),
            (
                "Review Letter",
                "Ringkasan review formal (opsional).",
                "Dapat diunduh/dilihat jika tersedia.",
                "Review Letter",
            ),
            (
                "Outcome",
                "Hasil review NOC.",
                "Reviewed/Rejected/Need Revision [TBD].",
                "[TBD]",
            ),
        ],
    )
    add_figure(doc, figures["noc_review_request"])

    doc.add_heading("5.4 Reports", level=2)
    add_paragraphs(
        doc,
        [
            "NOC dapat mengunduh laporan operasional untuk audit dan kebutuhan koordinasi lintas departemen.",
            "Catatan: Format, jenis laporan, dan penamaan file mengikuti standar internal [TBD].",
        ],
    )
    add_figure(doc, figures["noc_reports"])

    doc.add_heading("5.5 Recycle Material", level=2)
    add_paragraphs(
        doc,
        [
            "NOC memonitor laporan recycle dari OM dan memastikan tindak lanjut sesuai prosedur perusahaan.",
            "Aksi tindak lanjut (contoh): verifikasi, set status, dan pencatatan disposal [TBD].",
        ],
    )
    add_field_table(
        doc,
        "Elemen Data (Recycle Material - NOC View)",
        [
            ("Date", "Tanggal pelaporan.", "Read-only.", "23/04/2026"),
            ("Site/Reporter", "Site dan pelapor.", "Read-only.", "Papua / om"),
            ("Material", "Material dilaporkan.", "Read-only.", "[TBD]"),
            ("Condition", "Kondisi material.", "Read-only.", "Damaged"),
            ("Status", "Status tindak lanjut.", "Dapat diubah NOC [TBD].", "In Review"),
            ("Actions", "Aksi (detail/approve/reject) [TBD].", "Tergantung implementasi.", "[TBD]"),
        ],
    )
    add_figure(doc, figures["noc_recycle_material"])

    doc.add_heading("5.6 History Inventory", level=2)
    add_paragraphs(
        doc,
        [
            "NOC dapat memonitor riwayat pemakaian material oleh OM untuk analisis kebutuhan dan perencanaan replenishment.",
        ],
    )
    add_field_table(
        doc,
        "Elemen Data (History Inventory - NOC View)",
        [
            ("Tanggal", "Tanggal pemakaian.", "Read-only.", "23/04/2026"),
            ("Site", "Lokasi pemakaian.", "Read-only.", "Papua"),
            ("Material", "Nama material.", "Read-only.", "[TBD]"),
            ("Qty (pcs)", "Jumlah pemakaian.", "Read-only.", "5"),
            ("Project", "Proyek terkait.", "Read-only.", "Papua Solar"),
            ("Alasan", "Alasan pemakaian.", "Read-only.", "Pemeliharaan rutin"),
            ("Pelapor", "User pencatat.", "Read-only.", "om"),
        ],
    )
    add_figure(doc, figures["noc_history_inventory"])

    doc.add_heading("5.7 Shipping Control (NOC Shipping)", level=2)
    add_paragraphs(
        doc,
        [
            "Shipping Control adalah langkah birokrasi Step 4. Setelah GM approval, NOC memproses pengiriman ke site.",
            "Aktivitas utama:",
            "1. Menentukan ekspedisi dan input nomor resi.",
            "2. Menentukan ETA (perkiraan tanggal tiba).",
            "3. Mengunggah bukti pengiriman/foto jika tersedia [TBD].",
            "4. Memperbarui status pengiriman hingga delivered.",
        ],
    )
    add_field_table(
        doc,
        "Field Shipping (Contoh)",
        [
            ("Ekspedisi", "Nama jasa pengiriman.", "Diisi NOC.", "JNE"),
            ("Tracking/Resi No.", "Nomor resi pengiriman.", "Wajib.", "jne1245621331"),
            ("ETA", "Perkiraan tanggal tiba.", "Wajib/disarankan.", "04/27/2026"),
            ("Photo/Proof", "Bukti foto/label pengiriman.", "Opsional/disarankan.", "[TBD]"),
            ("Status Shipping", "Status pengiriman.", "Shipping/Delivered [TBD].", "Shipping"),
        ],
    )
    add_figure(doc, figures["noc_shipping_control"])

    doc.add_heading("5.8 Notifications", level=2)
    add_paragraphs(
        doc,
        [
            "Notifications pada NOC mencakup notifikasi permintaan material dan alert stok (warnings/critical).",
        ],
    )
    add_field_table(
        doc,
        "Kategori Notifikasi (NOC)",
        [
            (
                "Request Updates",
                "Perubahan status request dari OM/GM.",
                "Otomatis.",
                "Permintaan ... telah disetujui GM.",
            ),
            (
                "System Alerts",
                "Warning/Critical alert stok atau anomali.",
                "Otomatis.",
                "Stock ... mendekati batas",
            ),
            (
                "Shipping Updates",
                "Update pengiriman (resi/ETA/delivered).",
                "Otomatis/oleh NOC.",
                "Resi ...",
            ),
        ],
    )
    add_figure(doc, figures["noc_notifications"])
    add_figure(doc, figures["noc_notifications_monitor"])

    doc.add_heading("5.9 System Alerts", level=2)
    add_paragraphs(
        doc,
        [
            "System Alerts menampilkan alert otomatis terkait stok kritis dan anomali sistem. NOC bertanggung jawab untuk melakukan respon (resolve/ignore) sesuai prosedur.",
            "Catatan: Pengaturan threshold sumbernya dari katalog material (Stock Master).",
        ],
    )
    add_field_table(
        doc,
        "Elemen Data (System Alerts)",
        [
            ("Type", "Jenis alert.", "Warning/Critical.", "WARNING_STOCK"),
            ("Message", "Judul/deskripsi alert.", "Otomatis.", "Stock ... mendekati batas"),
            ("Lokasi", "Site yang terdampak.", "Otomatis.", "Papua"),
            ("Status", "Status alert.", "New/Read/Resolved [TBD].", "NEW"),
            ("Last Triggered", "Waktu terakhir alert muncul.", "Timestamp.", "23/04/2026 14:25"),
            ("Actions", "Respon alert.", "Ignore/Resolve Now.", "Resolve Now"),
            ("Catatan", "Catatan penanganan.", "Opsional.", "[TBD]"),
        ],
    )
    add_figure(doc, figures["noc_system_alerts"])

    doc.add_heading("5.10 Audit Logs", level=2)
    add_paragraphs(
        doc,
        [
            "Audit Logs mencatat jejak aktivitas transaksi dan perubahan database. Gunakan fitur Export CSV untuk kebutuhan audit.",
        ],
    )
    add_field_table(
        doc,
        "Kolom Audit Logs (Contoh)",
        [
            ("Waktu & Tanggal", "Timestamp aktivitas.", "Otomatis.", "23/04/2026 09:37"),
            ("User PIC", "User yang melakukan aksi.", "Otomatis.", "noc"),
            ("Aktivitas", "Jenis aksi.", "Create/Update/Delete/Approve [TBD].", "Update Stock"),
            ("Detail Modifikasi", "Ringkasan perubahan.", "Otomatis.", "[TBD]"),
            ("Modul", "Modul yang terdampak.", "Otomatis.", "Stock Master"),
        ],
    )
    add_figure(doc, figures["noc_audit_logs"])

    doc.add_heading("5.11 Settings", level=2)
    add_paragraphs(
        doc,
        [
            "Settings digunakan untuk pengelolaan profil dan keamanan akun NOC.",
        ],
    )
    add_field_table(
        doc,
        "Field Profile Akun (NOC)",
        [
            ("Username/Nama Lengkap", "Identitas akun.", "Tidak boleh kosong.", "noc"),
            ("Email Address", "Email user.", "Format email valid.", "noc@sundaya.co.id"),
            ("Nomor Telepon", "Kontak user.", "Format internasional disarankan.", "+62 812 3456 7890"),
            ("Bahasa/Language", "Preferensi bahasa UI.", "Dropdown.", "English"),
            ("Role Akun", "Role akses.", "Read-only (umumnya).", "NOC"),
        ],
    )
    add_figure(doc, figures["noc_settings"])

    doc.add_page_break()


def add_chapter_6_gm(doc: Document, figures: dict[str, Figure]) -> None:
    doc.add_heading("Bab 6 - GM (Executive / Final Approver)", level=1)
    add_paragraphs(
        doc,
        [
            "Role GM berfokus pada monitoring eksekutif dan pengambilan keputusan (final approval) atas permintaan material.",
            "Menu GM pada sidebar: Executive Monitor, Review Request, Reports, Recycle Material, History Inventory, Notifications, System Alerts, Audit Logs, Settings.",
        ],
    )

    doc.add_heading("6.1 Executive Monitor (Dashboard)", level=2)
    add_paragraphs(
        doc,
        [
            "Executive Monitor memberikan tampilan KPI dan analitik tingkat manajemen untuk seluruh site.",
            "Gunakan dashboard ini untuk melihat tren inbound/outbound, kebutuhan kritis, serta ringkasan alert sistem.",
        ],
    )
    add_field_table(
        doc,
        "Definisi KPI (Executive Monitor)",
        [
            ("Total Sites", "Jumlah site yang dimonitor.", "Tetap sesuai konfigurasi.", "3"),
            ("Total Stock Items", "Jumlah item pada katalog.", "Bergantung data master.", "[TBD]"),
            ("Critical Needs", "Indikator kebutuhan kritis.", "Berdasarkan alert/threshold.", "0"),
            ("Weekly Turnover", "Perubahan mingguan.", "Berdasarkan aktivitas 7 hari.", "+12%"),
            ("Alert Sistem Kritis", "Daftar alert prioritas.", "Menampilkan alert yang butuh tindakan.", "[TBD]"),
        ],
    )
    add_figure(doc, figures["gm_executive_monitor"])

    doc.add_heading("6.2 Review Request (GM Final Approval)", level=2)
    add_paragraphs(
        doc,
        [
            "Review Request pada GM adalah langkah birokrasi Step 3 (Final Approval).",
            "Alur keputusan:",
            "1. Buka request yang sudah direview oleh NOC.",
            "2. Tinjau urgensi, deadline, dan rekomendasi teknis (Review Letter jika tersedia).",
            "3. Putuskan approve atau reject sesuai kebijakan perusahaan.",
            "4. Sistem akan meneruskan ke NOC untuk proses shipping jika approved.",
        ],
    )
    add_field_table(
        doc,
        "Elemen Data (GM Final Approval)",
        [
            ("Request ID", "Nomor permintaan.", "Read-only.", "REQ-0020"),
            ("Urgency", "Prioritas permintaan.", "Read-only.", "Very Important"),
            ("Deadline", "Batas kebutuhan.", "Read-only.", "4/26/2026"),
            ("Review Letter", "Ringkasan review formal dari NOC (jika ada).", "Opsional.", "Review Letter"),
            ("Decision", "Keputusan GM.", "Approved/Rejected [TBD].", "[TBD]"),
            ("Rationale", "Alasan keputusan (jika diwajibkan).", "Opsional/Wajib [TBD].", "[TBD]"),
        ],
    )
    add_figure(doc, figures["gm_review_request"])
    add_figure(doc, figures["gm_review_letter_modal"])

    doc.add_heading("6.3 Reports", level=2)
    add_paragraphs(doc, ["GM dapat mengakses laporan sebagai bagian dari kontrol dan pelaporan manajemen."])
    doc.add_heading("6.3.1 Jenis Laporan (Ringkas)", level=3)
    rep = doc.add_table(rows=1, cols=3)
    rep.style = "Table Grid"
    rep.rows[0].cells[0].text = "Report"
    rep.rows[0].cells[1].text = "Tujuan"
    rep.rows[0].cells[2].text = "Output"
    for c in rep.rows[0].cells:
        c.paragraphs[0].runs[0].font.bold = True
    reports = [
        ("Stock Opname Report", "Kontrol stok per site.", "PDF"),
        ("Stock Mutation Report", "Monitoring mutasi stok.", "PDF"),
        ("Material Request Report", "Monitoring pipeline request.", "PDF"),
        ("Recent Movements Report", "Ringkasan aktivitas terbaru.", "PDF"),
    ]
    for r in reports:
        cells = rep.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v
    add_figure(doc, figures["gm_reports"])

    doc.add_heading("6.4 Recycle Material", level=2)
    add_paragraphs(doc, ["GM dapat memantau laporan recycle untuk memastikan kepatuhan proses dan efisiensi."])
    add_figure(doc, figures["gm_recycle_material"])

    doc.add_heading("6.5 History Inventory", level=2)
    add_paragraphs(doc, ["History Inventory membantu GM melihat pola pemakaian material lintas site dan proyek."])
    add_figure(doc, figures["gm_history_inventory"])

    doc.add_heading("6.6 Notifications", level=2)
    add_paragraphs(doc, ["Notifications menampilkan perubahan status request dan alert stok penting untuk monitoring manajemen."])
    add_field_table(
        doc,
        "Kategori Notifikasi (GM)",
        [
            ("Request Updates", "Perubahan status request yang perlu keputusan.", "Otomatis.", "Permintaan ... menunggu GM approval."),
            ("System Alerts", "Alert stok/anomali untuk awareness.", "Otomatis.", "Stock ... hampir habis"),
            ("Audit/Compliance", "Notifikasi terkait audit/kepatuhan [TBD].", "Opsional.", "[TBD]"),
        ],
    )
    add_figure(doc, figures["gm_notifications"])
    add_figure(doc, figures["gm_notifications_monitor"])

    doc.add_heading("6.7 System Alerts", level=2)
    add_paragraphs(doc, ["System Alerts pada GM berfungsi sebagai insight dan awareness terhadap risiko operasional."])
    add_field_table(
        doc,
        "Elemen Data (System Alerts - GM View)",
        [
            ("Type", "Jenis alert.", "Warning/Critical.", "CRITICAL_STOCK"),
            ("Lokasi", "Site terdampak.", "Otomatis.", "Maluku"),
            ("Status", "Status alert.", "New/Read/Resolved [TBD].", "READ"),
            ("Last Triggered", "Waktu muncul terakhir.", "Timestamp.", "23/04/2026 09:57"),
            ("Need Attention", "Indikator butuh tindakan.", "Berdasarkan severity.", "Yes"),
        ],
    )
    add_figure(doc, figures["gm_system_alerts"])

    doc.add_heading("6.8 Audit Logs", level=2)
    add_paragraphs(doc, ["Audit Logs dapat digunakan GM untuk memastikan kontrol internal dan jejak perubahan."])
    add_field_table(
        doc,
        "Kolom Audit Logs (GM View)",
        [
            ("Waktu & Tanggal", "Timestamp aktivitas.", "Otomatis.", "23/04/2026 09:37"),
            ("User PIC", "User pelaksana aksi.", "Otomatis.", "noc"),
            ("Aktivitas", "Jenis aksi.", "Otomatis.", "Approve Request"),
            ("Detail Modifikasi", "Ringkasan perubahan.", "Otomatis.", "[TBD]"),
            ("Modul", "Modul terdampak.", "Otomatis.", "Review Request"),
        ],
    )
    add_figure(doc, figures["gm_audit_logs"])

    doc.add_heading("6.9 Settings", level=2)
    add_paragraphs(doc, ["Settings untuk pengelolaan profil akun dan keamanan."])
    add_field_table(
        doc,
        "Field Profile Akun (GM)",
        [
            ("Username/Nama Lengkap", "Identitas akun.", "Tidak boleh kosong.", "gm_admin"),
            ("Email Address", "Email user.", "Format email valid.", "gm@sundaya.co.id"),
            ("Nomor Telepon", "Kontak user.", "Format internasional disarankan.", "+62 812 3456 7890"),
            ("Bahasa/Language", "Preferensi bahasa UI.", "Dropdown.", "English"),
            ("Role Akun", "Role akses.", "Read-only (umumnya).", "GM"),
        ],
    )
    add_figure(doc, figures["gm_settings"])

    doc.add_page_break()


def add_chapter_7_programmer(doc: Document, figures: dict[str, Figure]) -> None:
    doc.add_heading("Bab 7 - Programmer (System Administrator)", level=1)
    add_paragraphs(
        doc,
        [
            "Role Programmer memiliki akses penuh untuk monitoring sistem, konfigurasi, dan dukungan operasional. "
            "Perubahan pada flow, role, dan data master harus mengikuti prosedur change management internal.",
            "Menu Programmer pada sidebar: Programmer Control, Stock Master, Material Requests, Reports, Recycle Material, History Inventory, Shipping Control, Notifications, System Alerts, Audit Logs, Settings.",
        ],
    )

    doc.add_heading("7.1 Programmer Control", level=2)
    add_paragraphs(
        doc,
        [
            "Programmer Control Center menampilkan monitoring menyeluruh untuk alur operasional dan modul sistem.",
            "Contoh fungsi:",
            "1. Monitoring pipeline status request (menunggu NOC/GM, pengiriman, selesai).",
            "2. Kontrol distribusi role user dan akses cepat ke modul.",
            "3. Refresh data untuk sinkronisasi tampilan monitoring.",
        ],
    )
    add_field_table(
        doc,
        "Widget Monitoring (Programmer Control)",
        [
            ("Total Request", "Total permintaan pada pipeline.", "Bergantung filter/periode [TBD].", "0"),
            ("Menunggu NOC", "Permintaan menunggu review NOC.", "Otomatis.", "0"),
            ("Menunggu GM", "Permintaan menunggu approval GM.", "Otomatis.", "0"),
            ("Menunggu NOC Kirim", "Permintaan menunggu proses shipping NOC.", "Otomatis.", "0"),
            ("Dalam Pengiriman", "Permintaan sedang dikirim.", "Otomatis.", "0"),
            ("Selesai", "Permintaan selesai/fulfilled.", "Otomatis.", "0"),
            ("Total Site", "Jumlah site terdaftar.", "Konfigurasi sistem.", "0/3 [TBD]"),
            ("Notif Belum Dibaca", "Jumlah notifikasi unread.", "Otomatis.", "0"),
        ],
    )
    add_figure(doc, figures["programmer_control"])

    doc.add_heading("7.2 Stock Master", level=2)
    add_paragraphs(
        doc,
        [
            "Programmer memiliki akses pengelolaan katalog material dan threshold. "
            "Praktik terbaik: lakukan perubahan melalui prosedur persetujuan internal dan pastikan Audit Logs tercatat.",
            "Detail field Add/Edit Material pada Stock Master mengacu pada Bab 5.2 (NOC - Stock Master).",
        ],
    )
    add_figure(doc, figures["programmer_stock_master"])

    doc.add_heading("7.3 Material Requests", level=2)
    add_paragraphs(
        doc,
        [
            "Programmer dapat memonitor seluruh permintaan material untuk kebutuhan troubleshooting dan dukungan proses.",
        ],
    )
    add_field_table(
        doc,
        "Elemen Data (Material Requests - Programmer View)",
        [
            ("Request ID", "Nomor permintaan.", "Read-only.", "REQ-0020"),
            ("Status Step", "Posisi dalam alur 5 langkah.", "Otomatis.", "Menunggu NOC Review"),
            ("Tracking No.", "Resi pengiriman.", "Diisi NOC saat shipping.", "jne..."),
            ("ETA", "Perkiraan tanggal tiba.", "Diisi NOC.", "04/27/2026"),
            ("Receipt", "Konfirmasi penerimaan.", "Diisi OM.", "Items received"),
        ],
    )
    add_figure(doc, figures["programmer_material_requests"])

    doc.add_heading("7.4 Reports", level=2)
    add_paragraphs(doc, ["Reports untuk monitoring dan verifikasi data."])
    doc.add_heading("7.4.1 Jenis Laporan (Ringkas)", level=3)
    rep = doc.add_table(rows=1, cols=3)
    rep.style = "Table Grid"
    rep.rows[0].cells[0].text = "Report"
    rep.rows[0].cells[1].text = "Tujuan"
    rep.rows[0].cells[2].text = "Output"
    for c in rep.rows[0].cells:
        c.paragraphs[0].runs[0].font.bold = True
    reports = [
        ("Stock Opname Report", "Kontrol stok per site.", "PDF"),
        ("Stock Mutation Report", "Monitoring mutasi stok.", "PDF"),
        ("Material Request Report", "Monitoring pipeline request.", "PDF"),
        ("Recent Movements Report", "Ringkasan aktivitas terbaru.", "PDF"),
    ]
    for r in reports:
        cells = rep.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v
    add_figure(doc, figures["programmer_reports"])

    doc.add_heading("7.5 Recycle Material", level=2)
    add_paragraphs(doc, ["Monitoring laporan recycle serta dukungan proses tindak lanjut."])
    add_field_table(
        doc,
        "Elemen Data (Recycle Material - Programmer View)",
        [
            ("Date", "Tanggal pelaporan.", "Read-only.", "23/04/2026"),
            ("Site/Reporter", "Site dan pelapor.", "Read-only.", "Papua / om"),
            ("Material", "Material yang dilaporkan.", "Read-only.", "[TBD]"),
            ("Condition", "Kondisi material.", "Read-only.", "Damaged"),
            ("Status", "Status tindak lanjut.", "Dapat diubah sesuai hak akses [TBD].", "In Review"),
        ],
    )
    add_figure(doc, figures["programmer_recycle_material"])

    doc.add_heading("7.6 History Inventory", level=2)
    add_paragraphs(doc, ["Monitoring log pemakaian untuk kebutuhan audit dan analisis."])
    add_field_table(
        doc,
        "Elemen Data (History Inventory - Programmer View)",
        [
            ("Tanggal", "Tanggal pemakaian.", "Read-only.", "23/04/2026"),
            ("Site", "Lokasi pemakaian.", "Read-only.", "Papua"),
            ("Material", "Nama material.", "Read-only.", "[TBD]"),
            ("Qty (pcs)", "Jumlah pemakaian.", "Read-only.", "5"),
            ("Project", "Proyek terkait.", "Read-only.", "Papua Solar"),
            ("Alasan", "Alasan pemakaian.", "Read-only.", "Pemeliharaan rutin"),
            ("Pelapor", "User pencatat.", "Read-only.", "om"),
        ],
    )
    add_figure(doc, figures["programmer_history_inventory"])

    doc.add_heading("7.7 Shipping Control", level=2)
    add_paragraphs(doc, ["Shipping Control untuk monitoring pengiriman dan bukti serah terima."])
    add_field_table(
        doc,
        "Elemen Data (Shipping Control - Programmer View)",
        [
            ("Tracking/Resi No.", "Nomor resi.", "Read-only/diisi NOC.", "jne..."),
            ("ETA", "Perkiraan tiba.", "Read-only/diisi NOC.", "04/27/2026"),
            ("Status", "Status pengiriman.", "Otomatis.", "Shipping"),
            ("Proof", "Bukti foto/dokumen.", "Opsional.", "[TBD]"),
        ],
    )
    add_figure(doc, figures["programmer_shipping_control"])

    doc.add_heading("7.8 Notifications", level=2)
    add_paragraphs(doc, ["Notifikasi untuk monitoring status dan alert sistem."])
    add_field_table(
        doc,
        "Kategori Notifikasi (Programmer)",
        [
            ("Request Updates", "Perubahan status pipeline request.", "Otomatis.", "Menunggu GM"),
            ("System Alerts", "Warning/Critical alert stok/anomali.", "Otomatis.", "Critical stock"),
            ("Audit", "Notifikasi perubahan konfigurasi/role [TBD].", "Opsional.", "[TBD]"),
        ],
    )
    add_figure(doc, figures["programmer_notifications"])

    doc.add_heading("7.9 System Alerts", level=2)
    add_paragraphs(doc, ["System Alerts untuk investigasi alert stok/anomali dan verifikasi konfigurasi threshold."])
    add_field_table(
        doc,
        "Elemen Data (System Alerts - Programmer)",
        [
            ("Type", "Jenis alert.", "Warning/Critical.", "WARNING_STOCK"),
            ("Lokasi", "Site terdampak.", "Otomatis.", "Papua"),
            ("Status", "Status alert.", "New/Read/Resolved [TBD].", "NEW"),
            ("Actions", "Aksi respon.", "Ignore/Resolve Now.", "Resolve Now"),
        ],
    )
    add_figure(doc, figures["programmer_system_alerts"])

    doc.add_heading("7.10 Audit Logs", level=2)
    add_paragraphs(doc, ["Audit Logs sebagai sumber utama penelusuran perubahan dan troubleshooting."])
    add_field_table(
        doc,
        "Kolom Audit Logs (Programmer)",
        [
            ("Waktu & Tanggal", "Timestamp aktivitas.", "Otomatis.", "23/04/2026 09:37"),
            ("User PIC", "User pelaksana aksi.", "Otomatis.", "programmer"),
            ("Aktivitas", "Jenis aksi.", "Otomatis.", "Update Flow"),
            ("Detail Modifikasi", "Ringkasan perubahan.", "Otomatis.", "[TBD]"),
            ("Modul", "Modul terdampak.", "Otomatis.", "Settings"),
        ],
    )
    add_figure(doc, figures["programmer_audit_logs"])

    doc.add_heading("7.11 Settings", level=2)
    add_paragraphs(
        doc,
        [
            "Settings pada Programmer mencakup pengaturan profil, keamanan, serta kontrol tambahan seperti Admin Control dan Flow.",
            "Detail fungsi Admin Control & Flow (konfigurasi role/flow): [TBD] sesuai implementasi sistem.",
        ],
    )
    add_field_table(
        doc,
        "Field Profile Akun (Programmer)",
        [
            ("Username/Nama Lengkap", "Identitas akun.", "Tidak boleh kosong.", "programmer"),
            ("Email Address", "Email user.", "Format email valid.", "it@sundaya.co.id"),
            ("Nomor Telepon", "Kontak user.", "Format internasional disarankan.", "+62 812 3456 7890"),
            ("Bahasa/Language", "Preferensi bahasa UI.", "Dropdown.", "English"),
            ("Role Akun", "Role akses.", "Read-only (umumnya).", "PROGRAMMER"),
        ],
    )
    add_figure(doc, figures["programmer_settings"])

    doc.add_page_break()


def add_chapter_8_faq(doc: Document) -> None:
    doc.add_heading("Bab 8 - FAQ & Troubleshooting", level=1)

    doc.add_heading("8.1 FAQ (Pertanyaan Umum)", level=2)
    faqs = [
        ("Saya tidak bisa login. Apa yang harus dilakukan?", "Pastikan username/password benar, periksa koneksi internet, lalu coba login kembali. Jika masih gagal, lakukan reset password melalui admin [TBD]."),
        ("Menu saya tidak lengkap / berbeda dari rekan saya.", "Hak akses menu mengikuti role. Pastikan role akun Anda benar. Jika perlu perubahan role, ajukan ke Programmer melalui prosedur internal [TBD]."),
        ("Bagaimana cara mengetahui status request saya?", "Buka Material Requests (OM) atau Review Request (NOC/GM). Periksa panel Alur Birokrasi dan notifikasi perubahan status."),
        ("Kenapa muncul alert stok (warning/critical)?", "Alert muncul saat stok mendekati atau melewati threshold minimum. Verifikasi threshold di Stock Master dan rencanakan replenishment."),
        ("Bagaimana cara download laporan PDF?", "Buka menu Reports lalu klik Download PDF pada jenis laporan yang dibutuhkan. Pastikan pop-up download tidak diblokir browser."),
        ("Request saya stuck di status tertentu.", "Cek alur step mana yang belum selesai (NOC Review/GM Approval/NOC Shipping/OM Receiving). Jika melebihi SLA, lakukan eskalasi melalui kontak support."),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        p.add_run("Q: ").bold = True
        p.add_run(q)
        p2 = doc.add_paragraph()
        p2.add_run("A: ").bold = True
        p2.add_run(a)

    doc.add_heading("8.2 Troubleshooting", level=2)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    headers = ["Masalah", "Gejala", "Penyebab Umum", "Solusi Cepat"]
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    issues = [
        ("Gagal Login", "Pesan error login / tidak masuk dashboard", "Credential salah / akun nonaktif / koneksi bermasalah", "Coba ulang, reset password [TBD], hubungi support"),
        ("Halaman Blank", "UI kosong setelah klik menu", "Cache browser / session expired", "Reload, logout-login, clear cache"),
        ("Lambat / Timeout", "Proses loading lama", "Koneksi internet / beban server", "Cek koneksi, coba di jam berbeda, eskalasi ke Programmer"),
        ("Notifikasi Tidak Muncul", "Badge notifikasi tidak bertambah", "Permission browser / event belum terjadi", "Refresh, cek Notifications, verifikasi flow status"),
        ("Download Report Gagal", "PDF tidak terunduh", "Pop-up blocked / server generate report error", "Allow download, coba ulang, cek laporan lain, eskalasi"),
        ("Status Request Tidak Berubah", "Request stuck di step tertentu", "Pending action di role berikutnya", "Cek panel Alur Birokrasi, follow-up NOC/GM sesuai step"),
    ]
    for row in issues:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    doc.add_heading("8.3 Tabel Eskalasi Kontak Support", level=2)
    esc = doc.add_table(rows=1, cols=4)
    esc.style = "Table Grid"
    headers = ["Level", "Tim", "Kontak", "Kapan Digunakan"]
    for i, h in enumerate(headers):
        esc.rows[0].cells[i].text = h
        esc.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    rows = [
        ("L1", "NOC", "[TBD] (email/WA/telepon)", "Masalah operasional request, shipping, atau alert stok"),
        ("L2", "Programmer", "[TBD] (email/WA/telepon)", "Bug sistem, perubahan role/flow, error aplikasi"),
        ("L3", "GM", "[TBD] (email/telepon)", "Eskalasi persetujuan atau keputusan manajemen"),
    ]
    for row in rows:
        cells = esc.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build_figures() -> dict[str, Figure]:
    media = Path("docs/manual_source_extracted/word/media").resolve()

    def f(img: str, caption: str) -> Figure:
        return Figure(image_path=media / img, caption=caption)

    return {
        "login": f("image1.png", "Gambar 2-1. Halaman Login Sundaya Warehouse System"),
        # OM
        "om_site_overview": f("image12.png", "Gambar 4-1. OM - Site Overview (Pusat Dashboard)"),
        "om_stock_master": f("image11.png", "Gambar 4-2. OM - Stock Master (Stock & Catalog)"),
        "om_material_requests": f("image10.png", "Gambar 4-3. OM - Material Requests"),
        "om_new_request_modal": f("image9.png", "Gambar 4-4. OM - Form New Request (Material Request Form)"),
        "om_reports": f("image8.png", "Gambar 4-5. OM - Reports"),
        "om_recycle_material": f("image6.png", "Gambar 4-6. OM - Recycle Material"),
        "om_history_inventory": f("image5.png", "Gambar 4-7. OM - History Inventory"),
        "om_notifications": f("image4.png", "Gambar 4-8. OM - Notifications (Feed)"),
        "om_notifications_monitor": f("image3.png", "Gambar 4-9. OM - Monitor Notifikasi (Panel Ringkas)"),
        "om_settings": f("image2.png", "Gambar 4-10. OM - Settings"),
        # NOC
        "noc_control": f("image26.png", "Gambar 5-1. NOC - NOC Control Dashboard"),
        "noc_stock_master": f("image25.png", "Gambar 5-2. NOC - Stock Master"),
        "noc_add_material_modal": f("image24.png", "Gambar 5-3. NOC - Add New Material"),
        "noc_edit_material_modal": f("image23.png", "Gambar 5-4. NOC - Edit Material"),
        "noc_review_request": f("image22.png", "Gambar 5-5. NOC - Review Request (Technical Review)"),
        "noc_reports": f("image21.png", "Gambar 5-6. NOC - Reports"),
        "noc_recycle_material": f("image20.png", "Gambar 5-7. NOC - Recycle Material"),
        "noc_history_inventory": f("image19.png", "Gambar 5-8. NOC - History Inventory"),
        "noc_shipping_control": f("image18.png", "Gambar 5-9. NOC - Shipping Control"),
        "noc_notifications": f("image17.png", "Gambar 5-10. NOC - Notifications"),
        "noc_notifications_monitor": f("image13.png", "Gambar 5-11. NOC - Monitor Notifikasi (Panel Ringkas)"),
        "noc_system_alerts": f("image16.png", "Gambar 5-12. NOC - System Alerts"),
        "noc_audit_logs": f("image15.png", "Gambar 5-13. NOC - Audit Logs"),
        "noc_settings": f("image14.png", "Gambar 5-14. NOC - Settings"),
        # GM
        "gm_executive_monitor": f("image37.png", "Gambar 6-1. GM - Executive Monitor (Dashboard)"),
        "gm_review_request": f("image36.png", "Gambar 6-2. GM - Review Request (Final Approval)"),
        "gm_review_letter_modal": f("image35.png", "Gambar 6-3. GM - Request Review Letter"),
        "gm_reports": f("image33.png", "Gambar 6-4. GM - Reports"),
        "gm_recycle_material": f("image32.png", "Gambar 6-5. GM - Recycle Material"),
        "gm_history_inventory": f("image31.png", "Gambar 6-6. GM - History Inventory"),
        "gm_notifications": f("image30.png", "Gambar 6-7. GM - Notifications"),
        "gm_notifications_monitor": f("image34.png", "Gambar 6-8. GM - Monitor Notifikasi (Panel Ringkas)"),
        "gm_system_alerts": f("image29.png", "Gambar 6-9. GM - System Alerts"),
        "gm_audit_logs": f("image28.png", "Gambar 6-10. GM - Audit Logs"),
        "gm_settings": f("image27.png", "Gambar 6-11. GM - Settings"),
        # Programmer
        "programmer_control": f("image49.png", "Gambar 7-1. Programmer - Control Center"),
        "programmer_stock_master": f("image48.png", "Gambar 7-2. Programmer - Stock Master"),
        "programmer_material_requests": f("image46.png", "Gambar 7-3. Programmer - Material Requests"),
        "programmer_reports": f("image45.png", "Gambar 7-4. Programmer - Reports"),
        "programmer_recycle_material": f("image44.png", "Gambar 7-5. Programmer - Recycle Material"),
        "programmer_history_inventory": f("image43.png", "Gambar 7-6. Programmer - History Inventory"),
        "programmer_shipping_control": f("image42.png", "Gambar 7-7. Programmer - Shipping Control"),
        "programmer_notifications": f("image41.png", "Gambar 7-8. Programmer - Notifications"),
        "programmer_system_alerts": f("image40.png", "Gambar 7-9. Programmer - System Alerts"),
        "programmer_audit_logs": f("image39.png", "Gambar 7-10. Programmer - Audit Logs"),
        "programmer_settings": f("image38.png", "Gambar 7-11. Programmer - Settings"),
    }


def main() -> None:
    figures = build_figures()
    logo_path = Path("docs/assets/sundaya_logo.png").resolve()
    out = Path("docs/Manual Book - Sundaya Warehouse System V2.0.docx").resolve()

    doc = Document()
    set_default_styles(doc)
    set_page_layout(doc)

    add_cover(doc, logo_path=logo_path)
    add_document_control(doc)
    add_toc(doc)
    doc.add_page_break()

    add_doc_conventions(doc)
    add_chapter_1(doc)
    add_chapter_2(doc, figures=figures)
    add_chapter_3(doc)
    add_chapter_4_om(doc, figures=figures)
    add_chapter_5_noc(doc, figures=figures)
    add_chapter_6_gm(doc, figures=figures)
    add_chapter_7_programmer(doc, figures=figures)
    add_chapter_8_faq(doc)

    doc.save(str(out))
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
